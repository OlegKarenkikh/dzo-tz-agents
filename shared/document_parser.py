"""
shared/document_parser.py
Structured data extraction from participant anketa files (PDF and DOCX).

Parses the standard 15-field anketa table and returns an ``AnketaData``
dataclass. Supports both PDF (via pdfplumber) and DOCX (via python-docx).

Usage::

    from shared.document_parser import parse_anketa

    data = parse_anketa(file_bytes, "application/pdf")
    print(data.company_name, data.inn)
"""
from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field

from shared.logger import setup_logger  # noqa: F401

logger = logging.getLogger("document_parser")


@dataclass
class AnketaData:
    """Structured data extracted from a participant anketa."""
    tender_id: str = ""
    company_name: str = ""
    org_form: str = ""
    inn: str = ""
    kpp: str = ""
    ogrn: str = ""
    legal_address: str = ""
    actual_address: str = ""
    postal_address: str = ""
    phone: str = ""
    email: str = ""
    website: str = ""
    bank_details: str = ""
    holding_membership: str = ""
    tax_system: str = ""
    subcontractors: str = ""
    authorized_person: str = ""
    responsible_contact: str = ""
    raw_fields: dict[int, str] = field(default_factory=dict)


@dataclass
class NDAData:
    """Extracted info from NDA documents."""
    signatory_name: str = ""
    signatory_position: str = ""
    company_name: str = ""
    date: str = ""


# Row number (1-based) → AnketaData field name mapping
_FIELD_MAP: dict[int, str] = {
    1: "company_name",
    2: "org_form",
    3: "inn",  # May contain INN, KPP, OGRN combined
    4: "legal_address",
    5: "actual_address",
    6: "postal_address",
    7: "phone",
    8: "email",
    9: "website",
    10: "bank_details",
    11: "holding_membership",
    12: "tax_system",
    13: "subcontractors",
    14: "authorized_person",
    15: "responsible_contact",
}

# Pattern to extract tender ID from anketa header
_TENDER_ID_RE = re.compile(
    r"(?:АНКЕТА\s+УЧАСТНИКА\s+(?:ТЕНДЕРНОГО\s+ОТБОРА|ТО)\s+)"
    r"(\d{3,5}[-–—]\s*[А-Яа-яA-Za-z]{2,}[-–—]\s*[А-Яа-яA-Za-z0-9]+)",
    re.IGNORECASE,
)

# Fallback: just look for the ID pattern anywhere
_TENDER_ID_FALLBACK_RE = re.compile(
    r"\d{3,5}[-–—]\s*[А-Яа-яA-Za-z]{2,}[-–—]\s*[А-Яа-яA-Za-z0-9]+",
)

# INN pattern (10 or 12 digits)
_INN_RE = re.compile(r"\b(\d{10}(?:\d{2})?)\b")
_KPP_RE = re.compile(r"КПП[:\s]*(\d{9})", re.IGNORECASE)
_OGRN_RE = re.compile(r"ОГРН[:\s]*(\d{13,15})", re.IGNORECASE)


def parse_anketa(file_bytes: bytes, content_type: str = "") -> AnketaData:
    """Parse anketa file and extract all 15 fields.

    Args:
        file_bytes: Raw file content
        content_type: MIME type or file extension hint

    Returns:
        AnketaData with extracted fields
    """
    ct = content_type.lower()
    if "pdf" in ct or ct.endswith(".pdf"):
        return _parse_pdf(file_bytes)
    elif "docx" in ct or "word" in ct or ct.endswith(".docx"):
        return _parse_docx(file_bytes)
    elif "doc" in ct:
        # Try DOCX first, fall back to PDF
        try:
            return _parse_docx(file_bytes)
        except Exception:
            return _parse_pdf(file_bytes)
    # Default: try DOCX then PDF
    try:
        return _parse_docx(file_bytes)
    except Exception:
        return _parse_pdf(file_bytes)


def parse_nda(file_bytes: bytes, content_type: str = "") -> NDAData:
    """Parse NDA document to extract signatory information."""
    ct = content_type.lower()
    text = ""
    if "pdf" in ct:
        text = _extract_pdf_text(file_bytes)
    elif "docx" in ct or "word" in ct:
        text = _extract_docx_text(file_bytes)
    else:
        try:
            text = _extract_docx_text(file_bytes)
        except Exception:
            text = _extract_pdf_text(file_bytes)

    return _parse_nda_text(text)


# ---------------------------------------------------------------------------
#  DOCX parsing
# ---------------------------------------------------------------------------

def _parse_docx(file_bytes: bytes) -> AnketaData:
    """Extract anketa data from a DOCX file."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    data = AnketaData()

    # Extract tender ID from paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        m = _TENDER_ID_RE.search(text)
        if m:
            data.tender_id = m.group(1).strip()
            break
        m = _TENDER_ID_FALLBACK_RE.search(text)
        if m and not data.tender_id:
            data.tender_id = m.group(0).strip()

    # Find the main anketa table (first table with >= 15 rows)
    for table in doc.tables:
        if len(table.rows) < 3:
            continue
        _extract_from_docx_table(table, data)
        if data.company_name:
            break

    return data


def _extract_from_docx_table(table, data: AnketaData) -> None:
    """Extract fields from a DOCX table."""
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3:
            continue

        # Try to get row number from first cell
        try:
            row_num = int(cells[0])
        except (ValueError, IndexError):
            continue

        # The value is in the last non-empty cell (typically column 3)
        value = cells[-1].strip()
        if not value:
            # Try second-to-last
            for c in reversed(cells[1:]):
                if c.strip() and not c.strip().isdigit():
                    value = c.strip()
                    break

        data.raw_fields[row_num] = value
        field_name = _FIELD_MAP.get(row_num)
        if field_name and value:
            if field_name == "inn":
                _parse_inn_field(value, data)
            else:
                setattr(data, field_name, value)


# ---------------------------------------------------------------------------
#  PDF parsing
# ---------------------------------------------------------------------------

def _parse_pdf(file_bytes: bytes) -> AnketaData:
    """Extract anketa data from a PDF file."""
    import pdfplumber

    data = AnketaData()

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        # Extract text for tender ID
        full_text = ""
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            full_text += page_text + "\n"

        # Extract tender ID
        m = _TENDER_ID_RE.search(full_text)
        if m:
            data.tender_id = m.group(1).strip()
        else:
            m = _TENDER_ID_FALLBACK_RE.search(full_text)
            if m:
                data.tender_id = m.group(0).strip()

        # Extract tables
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                _extract_from_pdf_table(table, data)
                if data.company_name:
                    break
            if data.company_name:
                break

    return data


def _extract_from_pdf_table(table: list[list], data: AnketaData) -> None:
    """Extract fields from a pdfplumber table."""
    for row in table:
        if not row or len(row) < 2:
            continue

        # Clean cells: remove None, strip whitespace
        cells = [str(c or "").strip() for c in row]

        # Try to get row number from first cell
        try:
            row_num = int(cells[0])
        except (ValueError, IndexError):
            continue

        # Find the value cell — skip the number and label columns
        # PDF tables may have merged cells resulting in None/empty
        value = ""
        for c in reversed(cells[1:]):
            c_clean = c.strip().replace("\n", " ")
            # Skip if it looks like a label (long text with no data)
            if c_clean and not _is_label_text(c_clean, row_num):
                value = c_clean
                break

        if not value:
            # Try to get from the last non-empty cell
            for c in reversed(cells):
                c_clean = c.strip()
                if c_clean and c_clean != str(row_num):
                    value = c_clean.replace("\n", " ")
                    break

        data.raw_fields[row_num] = value
        field_name = _FIELD_MAP.get(row_num)
        if field_name and value:
            if field_name == "inn":
                _parse_inn_field(value, data)
            else:
                setattr(data, field_name, value)


def _is_label_text(text: str, row_num: int) -> bool:
    """Check if text looks like a field label rather than a value."""
    labels = {
        1: "наименование",
        2: "организационно-правовая",
        3: "инн",
        4: "юридический",
        5: "фактический",
        6: "почтовый",
        7: "телефон",
        8: "адрес электронной",
        9: "адрес сайта",
        10: "банковские",
        11: "вхождение",
        12: "сведения об отнесении",
        13: "сведения о привлечении",
        14: "ф.и.о. уполномоченного",
        15: "ф.и.о. лица",
    }
    label = labels.get(row_num, "")
    if label and text.lower().startswith(label):
        return True
    return False


# ---------------------------------------------------------------------------
#  Helpers
# ---------------------------------------------------------------------------

def _parse_inn_field(value: str, data: AnketaData) -> None:
    """Parse the INN/KPP/OGRN combined field."""
    # Extract INN
    inn_match = _INN_RE.search(value)
    if inn_match:
        data.inn = inn_match.group(1)

    # Extract KPP
    kpp_match = _KPP_RE.search(value)
    if kpp_match:
        data.kpp = kpp_match.group(1)

    # Extract OGRN
    ogrn_match = _OGRN_RE.search(value)
    if ogrn_match:
        data.ogrn = ogrn_match.group(1)

    # If no separate INN found, use entire value as INN
    if not data.inn and value.strip().isdigit():
        data.inn = value.strip()


def _extract_pdf_text(file_bytes: bytes) -> str:
    """Extract full text from PDF."""
    import pdfplumber
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        return "\n".join(p.extract_text() or "" for p in pdf.pages)


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract full text from DOCX."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


# ---------------------------------------------------------------------------
#  TZ Section Detection (rules-engine)
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(
    r"^(?:"
    r"\d+\.(?:\d+\.)*\s+"       # numbered: "1. ", "2.1. "
    r"|[IVXLC]+\.\s+"           # roman: "II. "
    r"|#{1,4}\s+"               # markdown: "## "
    r"|[А-ЯA-Z\s]{4,}$"        # ALL CAPS line (min 4 chars)
    r")",
    re.MULTILINE,
)


def _is_heading_line(line: str) -> bool:
    """Check if a line looks like a section heading."""
    stripped = line.strip()
    if not stripped:
        return False
    # Numbered heading: "1. Title", "2.1. Title"
    if re.match(r"^\d+\.(?:\d+\.)*\s+\S", stripped):
        return True
    # Roman numeral heading
    if re.match(r"^[IVXLC]+\.\s+\S", stripped):
        return True
    # Markdown heading
    if stripped.startswith("#"):
        return True
    # ALL CAPS line (at least 4 chars, mostly uppercase letters)
    alpha_chars = [c for c in stripped if c.isalpha()]
    if len(alpha_chars) >= 4 and sum(1 for c in alpha_chars if c.isupper()) / len(alpha_chars) > 0.7:
        return True
    return False


def _find_in_headings(text: str, keywords: list[str]) -> bool:
    """Check if any keyword appears in a heading line."""
    for line in text.split("\n"):
        if not _is_heading_line(line):
            continue
        line_lower = line.lower()
        for kw in keywords:
            if kw.lower() in line_lower:
                return True
    return False


def _find_in_context(text: str, keyword: str, context_keywords: list[str], max_distance: int = 100) -> bool:
    """Check if keyword appears near context keywords (within max_distance chars)."""
    text_lower = text.lower()
    kw_lower = keyword.lower()
    pos = 0
    while True:
        idx = text_lower.find(kw_lower, pos)
        if idx == -1:
            break
        context_window = text_lower[max(0, idx - max_distance):idx + len(kw_lower) + max_distance]
        for ckw in context_keywords:
            if ckw.lower() in context_window:
                return True
        pos = idx + 1
    return False


def detect_sections(text: str) -> dict[str, bool]:
    """Detect presence of standard TZ sections using structural analysis.

    Unlike naive keyword matching, this function checks that keywords
    appear in section HEADINGS or in appropriate CONTEXT, not just
    anywhere in the document body.

    Returns dict with boolean flags for each standard section.
    """
    result = {}

    # has_goal: "Цель закупки" must appear as heading or in structured context
    result["has_goal"] = _find_in_headings(text, [
        "цель закупки", "цель ", "общие положения",
    ]) or _find_in_context(text, "целью", ["закупк", "обеспечен"], 150)

    # has_requirements: technical requirements section
    result["has_requirements"] = _find_in_headings(text, [
        "требовани", "технические характеристики", "спецификаци",
    ]) or _find_in_headings(text, ["параметр"])

    # has_quantities: quantities with units — look for "N шт.", "N ед.", etc.
    qty_pattern = re.compile(
        r"\b\d+\s*(?:шт|ед|компл|комплект|лицензи|рулон|мешок|банк|пачек|пачк|коробок|штук)\b",
        re.IGNORECASE,
    )
    result["has_quantities"] = bool(qty_pattern.search(text))

    # has_delivery_term: delivery timeline
    term_patterns = [
        r"\b(?:срок(?:ам)?\s+(?:поставки|выполнения|оказания|передачи))",
        r"\b(?:в\s+течение\s+\d+\s*(?:\([^)]*\)\s*)?(?:календарных|рабочих)?\s*дней)",
        r"\b(?:не\s+позднее\s+\d+\s*(?:\([^)]*\)\s*)?дней)",
        r"\b\d+\s*(?:\([^)]*\)\s*)?(?:календарных|рабочих)\s+(?:дней|месяцев)",
    ]
    result["has_delivery_term"] = any(
        re.search(p, text, re.IGNORECASE) for p in term_patterns
    )

    # has_delivery_address: actual delivery address (not just "г." in specs)
    # Must find address-like context: street + building, or "место поставки" heading
    address_in_heading = _find_in_headings(text, ["место поставки", "место доставки", "адрес поставки"])
    address_pattern = re.compile(
        r"(?:ул\.|улица|пр\.|проспект|пер\.|переулок|ш\.|шоссе|наб\.|набережная)"
        r"[^.]{2,50}"
        r"(?:д\.|дом|стр\.|строение|корп\.|корпус)",
        re.IGNORECASE,
    )
    # "место поставки" or "адрес доставки" followed by actual address
    place_keyword = re.search(
        r"(?:место\s+поставки|место\s+оказания|адрес\s+доставки)[:\s]+[^\n]{5,}",
        text,
        re.IGNORECASE,
    )
    result["has_delivery_address"] = bool(
        address_in_heading or address_pattern.search(text) or place_keyword
    )

    # has_regulatory_reference: GOST, TR TS, Federal Law references
    # Must be standalone words, not inside serial numbers
    regulatory_patterns = [
        r"\bГОСТ\s+(?:Р\s+)?[\d.]",
        r"\bТР\s+ТС\b",
        r"\b(?:44|223)-ФЗ\b",
        r"\bФедеральн\w+\s+закон\w*\b",
        r"\bтехнич\w+\s+регламент\w*\b",
        r"\bИСО\s+\d",
        r"\bISO\s+\d",
    ]
    result["has_regulatory_reference"] = any(
        re.search(p, text, re.IGNORECASE) for p in regulatory_patterns
    )

    # has_evaluation_criteria: evaluation/scoring criteria
    criteria_in_heading = _find_in_headings(text, [
        "критерии оценки", "критерии выбора", "оценка предложений",
        "порядок оценки",
    ])
    criteria_pattern = re.compile(
        r"(?:цена|стоимость|квалификация|опыт|качество)\s*[—–-]\s*\d+\s*%",
        re.IGNORECASE,
    )
    result["has_evaluation_criteria"] = bool(
        criteria_in_heading or criteria_pattern.search(text)
    )

    return result


def compute_tz_score(sections: dict[str, bool]) -> float:
    """Compute TZ completeness score based on section weights."""
    weights = {
        "has_goal": 15.0,
        "has_requirements": 20.0,
        "has_quantities": 10.0,
        "has_delivery_term": 10.0,
        "has_delivery_address": 10.0,
        "has_regulatory_reference": 10.0,
        "has_evaluation_criteria": 10.0,
    }
    # has_requirements gets extra weight
    total_weight = sum(weights.values())
    achieved = sum(w for k, w in weights.items() if sections.get(k, False))
    return round(achieved / total_weight * 100, 1)


def compute_tz_decision(sections: dict[str, bool], score: float) -> str:
    """Compute rules-engine decision for TZ document."""
    critical_missing = []
    if not sections.get("has_goal"):
        critical_missing.append("цель закупки")
    if not sections.get("has_requirements"):
        critical_missing.append("технические требования")
    if not sections.get("has_quantities"):
        critical_missing.append("количество")

    if len(critical_missing) >= 2 or score < 50:
        return "ВЕРНУТЬ НА ДОРАБОТКУ"
    elif len(critical_missing) == 1 or score < 85:
        return "ПРИНЯТЬ С ЗАМЕЧАНИЕМ"
    else:
        return "ПРИНЯТЬ"


def _parse_nda_text(text: str) -> NDAData:
    """Parse NDA text to extract signatory info."""
    data = NDAData()

    # Look for signatory patterns
    signatory_re = re.compile(
        r"(?:подписант|уполномоченное\s+лицо|от\s+имени)[:\s]+(.+?)(?:\n|$)",
        re.IGNORECASE,
    )
    m = signatory_re.search(text)
    if m:
        data.signatory_name = m.group(1).strip()

    # Look for company name
    company_re = re.compile(
        r"(?:ООО|АО|ЗАО|ПАО|СПАО)\s*[«\"](.+?)[»\"]",
        re.IGNORECASE,
    )
    m = company_re.search(text)
    if m:
        data.company_name = m.group(0).strip()

    # Look for date
    date_re = re.compile(r"\d{1,2}[.\s]+\w+[.\s]+\d{4}")
    m = date_re.search(text)
    if m:
        data.date = m.group(0).strip()

    return data
